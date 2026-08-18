#!/usr/bin/env python3
"""Debug: read Crestmead Dental docx to see its actual table structure."""
from docx import Document
from pathlib import Path

P = Path(r"C:\Users\hamma\Downloads\New folder\Case Studies\Case Studies\Crestmead Dental\Crestmead_Dental_SEO_Case_Study_Omni_Path_Marketing.docx")
doc = Document(str(P))
print(f"=== {P.name} ===")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}\n")

for i, t in enumerate(doc.tables):
    print(f"--- Table {i}: {len(t.rows)}x{len(t.columns)} ---")
    for r, row in enumerate(t.rows[:6]):
        cells = [c.text.strip().replace('\n', ' | ')[:50] for c in row.cells]
        print(f"  R{r}: {cells}")
    print()
