#!/usr/bin/env python3
"""Read Marham.pk case study to see if format differs (non-dental, non-AU)."""
from docx import Document
from pathlib import Path

P = Path(r"C:\Users\hamma\Downloads\New folder\Case Studies\Case Studies\Marham.pk\Marham_PK_SEO_Case_Study_Omni_Path_Marketing.docx")
doc = Document(str(P))
print(f"=== {P.name} ===\n")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}\n")

print("--- Headings + first 2 paragraphs under each ---")
heading = None
for p in doc.paragraphs:
    s = p.style.name if p.style else "Normal"
    if s.startswith("Heading"):
        heading = p.text
        print(f"\n## {p.text}")
    elif p.text.strip() and heading:
        print(f"  {p.text[:160]}")
        heading = None  # only show first paragraph after each heading

print("\n--- Table 0 (top stats) ---")
t = doc.tables[0]
for row in t.rows:
    cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
    print(f"  {cells}")

print("\n--- Table 1 (intent breakdown) ---")
t = doc.tables[1]
for row in t.rows[:8]:
    cells = [c.text.strip()[:30] for c in row.cells]
    print(f"  {cells}")

print(f"\n--- Tables 2 and 3 sizes: {len(doc.tables[2].rows)}x{len(doc.tables[2].columns)} and {len(doc.tables[3].rows)}x{len(doc.tables[3].columns)} ---")
