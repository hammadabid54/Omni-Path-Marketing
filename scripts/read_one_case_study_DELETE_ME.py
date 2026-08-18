#!/usr/bin/env python3
"""Read the first Omni Path version case study to understand format."""
from docx import Document
from pathlib import Path

P = Path(r"C:\Users\hamma\Downloads\New folder\Case Studies\Case Studies\Bella Dental\Bella_Dental_SEO_Case_Study_Omni_Path_Marketing.docx")

doc = Document(str(P))
print(f"=== {P.name} ===\n")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print()

print("--- First 30 paragraphs ---")
for i, p in enumerate(doc.paragraphs[:30]):
    style = p.style.name if p.style else "Normal"
    text = p.text.strip()
    if text:
        print(f"[{style}] {text[:200]}")
    else:
        print(f"[{style}] (empty)")

print("\n--- Tables ---")
for i, t in enumerate(doc.tables):
    print(f"\nTable {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    for r, row in enumerate(t.rows[:6]):
        cells = [c.text.strip()[:40] for c in row.cells]
        print(f"  R{r}: {cells}")
    if len(t.rows) > 6:
        print(f"  ... {len(t.rows) - 6} more rows")
